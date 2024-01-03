# -*- coding: utf-8 -*-
# @Time    : 2024/1/2 22:11
# @Author  : Zhang Huan
# @Email   : johnhuan@whu.edu.cn
# QQ       : 248404941
# @File    : main.py

from docx import Document
from datetime import datetime as dt, timedelta
from datetime import date as dt_date
from zhdate import datetime as zh_dt
from zhdate import ZhDate as zh_date


def generate_school_calendar(year, month, day, week_list):
    for i in range(0, 20):
        week_name = "第%d周" % (i + 1)
        document.add_heading(week_name, 2)
        table = document.add_table(rows=4, cols=7, style='Colorful List')
        for j in range(0, 7):
            days = i * 7 + j
            date = dt_date(year, month, day) + timedelta(days=days)
            week_day = week_list[date.isoweekday() - 1]
            hdr_cells_mo_dy = table.rows[0].cells
            hdr_cells_mo_dy[j].text = "%d月%d" % (date.month, date.day)
            hdr_cells_zh_dy = table.rows[1].cells
            zh_date.from_datetime(dt(year, month, day))
            zh_d = zh_date.from_datetime(dt(date.year, date.month, date.day)).chinese()[5:-8]
            hdr_cells_zh_dy[j].text = zh_d
            hdr_cells_week_dy = table.rows[2].cells
            hdr_cells_week_dy[j].text = week_day
            hdr_cells_work = table.rows[3].cells
            hdr_cells_work[j].text = ""


if __name__ == '__main__':
    document = Document()
    document.add_heading("工作计划", 0)
    """
    第一学期
    """
    week_list = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    document.add_heading("2023~2024学年第一学期", 1)
    year = 2023
    month = 9
    day = 4
    generate_school_calendar(year, month, day, week_list)

    """
    第二学期
    """
    week_list = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    document.add_heading("2023~2024学年第二学期", 1)
    year = 2024
    month = 2
    day = 26
    generate_school_calendar(year, month, day, week_list)

    document.save("2023~2024年度工作计划.docx")
